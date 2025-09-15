from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Maak nieuw document aan
doc = Document()

# Voeg titel toe met naam
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
run = title.add_run("Jason Meulemans")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(255, 255, 255)

# Voeg achtergrondkleur toe aan de titelregel
section = doc.sections[0]
header = section.header
header_paragraph = header.paragraphs[0]
header_paragraph.text = " "
header_paragraph.style.font.size = Pt(1)
header_paragraph.paragraph_format.space_after = Pt(0)

# Voeg contactgegevens toe
contact = doc.add_paragraph()
contact.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
contact_run = contact.add_run("📧 jasonmeulemans@hotmail.com | 📞 +32 491 25 39 03 | 📍 Mechelen")
contact_run.font.size = Pt(10)



# Voeg sectie "Opleidingen" toe
doc.add_heading("Opleidingen", level=2)
educations = [
    ("Programmeren (Graduaat)", "Artesis Plantijn, Antwerpen", "sep 2024 – heden"),
    ("Electronica & ICT (Bachelor)", "Artesis Plantijn, Antwerpen", "sep 2022 – jul 2024"),
    ("Internet of Things (Graduaat)", "Artesis Plantijn, Antwerpen", "sep 2020 – jul 2022"),
    ("Digital Designer", "Crescendo CVO, Mechelen", "sep 2019 – jul 2020"),
    ("Carrosserie (Plaatbewerking)", "TSM, Mechelen", "sep 2014 – jul 2019")
]
for title, location, period in educations:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    p.add_run(f"\n{location}\n{period}")

# Voeg werkervaring toe
doc.add_heading("Werkervaring", level=2)
jobs = [
    ("OrderPicker", "Reynaers Aluminium, Duffel", "jan 2024 - heden",
     "Profielen van 7 Meter sorteren, ophalen, inpakken per klant/bedrijf"),
    ("Automatisation Engineer", "AAQUA Waterverzuivering, Mechelen", "jan 2020 - dec 2021",
     "NodeRed Flows configureren, Dataverwerking voor UI, Handleidingen maken voor UI-collega's/bedrijven, oplossingen voor automatische verwerking van Excel-data per machine."),
    ("Winkelbediende", "VapeShop, Mechelen", "jul 2019 - sep 2020",
     "Administratie van stock en productenlijst, klanten bedienen, testen van product, onderhoud en reparatie van gekochte producten"),
    ("Mechanieker", "Autopartners Renault, Mechelen", "jan 2017 - dec 2018",
     "Onderhoud, schuren en plamuren, montage, herstellen, matteren van voertuigen en stock management"),
    ("Mechanieker", "Seat Garage Hendrix, Mechelen", "jan 2016 - dec 2017",
     "Onderhoud en montage van voertuigen en matteren van onderdelen"),
    ("Mechanieker", "NMBS, Mechelen", "jan 2015 - jan 2016",
     "Onderhoud, schilderen, schuren en plamuren van wagons"),
    ("Onderhoud", "Stad Mechelen", "jan 2014 - jan 2015",
     "Montage en onderhoud van voertuigen en machines")
]
for title, company, period, desc in jobs:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    p.add_run(f"\n{company}\n{period}\n{desc}")

# Vaardigheden
doc.add_heading("Vaardigheden", level=2)
doc.add_paragraph("Programmeertalen: HTML, CSS, JavaScript, Node.js, MySQL, C#, Bash")
doc.add_paragraph("Tools: Linux, Cisco, Office 365")

# Talen
doc.add_heading("Talen", level=2)
doc.add_paragraph("Nederlands (moedertaal), Engels (goed)")

# Certificaten
doc.add_heading("Certificaten", level=2)
certs = [
    "Deelcertificaat: Websiteproductie",
    "Deelcertificaat: Grafische Vormgeving",
    "Deelcertificaat: Web animatie",
    "Deelcertificaat: Interfacedesign",
    "Deelcertificaat: Project Webdesign",
    "Certificaat: Websiteproductie",
    "Deelcertificaat: Eenvoudige Content Aanmaken",
    "Deelcertificaat: Digitale Beeldverwerking"
]
for cert in certs:
    doc.add_paragraph(cert, style='List Bullet')

# Eigenschappen
doc.add_heading("Eigenschappen", level=2)
traits = ["Teamplayer", "Leergierig", "Klantvriendelijk", "Ambitieus", "Zelfstandig", "Oplossingsgericht"]
for trait in traits:
    doc.add_paragraph(trait, style='List Bullet')

# Bestand opslaan

output_path = "C:/Users/jason/OneDrive/Desktop/01-PROGRAMMEREN AP 24-25/Sem-2/WPLexploratie/solliciteren/CV_Jason_Meulemans_Professioneel_2025.docx"
doc.save(output_path)

output_path